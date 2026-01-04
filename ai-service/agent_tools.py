import pandas as pd
from openpyxl import Workbook
from docx import Document
import uuid
import os
import shutil

# Configuration
# The AI Service runs on port 5002. We serve files from the /static endpoint.
BASE_URL = "http://localhost:5002" 
EXPORT_DIR = "static/generated_docs"

# Ensure the export directory exists
os.makedirs(EXPORT_DIR, exist_ok=True)

def save_file_locally(source_filename):
    """
    Moves the generated file to the static folder and returns a download URL.
    """
    destination_path = os.path.join(EXPORT_DIR, source_filename)
    
    # Move file from root to static folder
    if os.path.exists(source_filename):
        shutil.move(source_filename, destination_path)
    
    # Return the public URL
    return f"{BASE_URL}/{EXPORT_DIR}/{source_filename}"

def generate_roadmap_excel(steps):
    """
    Generates an Excel file using openpyxl.
    """
    filename = f"roadmap_{uuid.uuid4()}.xlsx"
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Plan"

        # Check if we have data
        if isinstance(steps, list) and len(steps) > 0:
            # If it's a list of dictionaries, use keys as headers
            if isinstance(steps[0], dict):
                headers = list(steps[0].keys())
                # Add Headers (Capitalized)
                ws.append([h.capitalize() for h in headers]) 
                
                # Add Rows
                for step in steps:
                    ws.append([str(step.get(h, "")) for h in headers])
            
            # If it's a simple list of strings
            elif isinstance(steps[0], str):
                ws.append(["Items"])
                for step in steps:
                    ws.append([step])

        wb.save(filename)
        
        # Save locally and get URL
        public_url = save_file_locally(filename)
        return public_url
    except Exception:
        return None

def generate_roadmap_word(content_summary, steps):
    """
    Generates a Word Document with a summary and a table of steps.
    """
    filename = f"roadmap_{uuid.uuid4()}.docx"
    doc = Document()
    
    doc.add_heading('Project Plan / Roadmap', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(content_summary)
    
    doc.add_heading('Details', level=1)
    
    # Create table if steps are structured
    if isinstance(steps, list) and len(steps) > 0 and isinstance(steps[0], dict):
        headers = list(steps[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header.capitalize()
        
        # Data Rows
        for step in steps:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(step.get(header, ''))
    
    doc.save(filename)
    
    # Save locally and get URL
    public_url = save_file_locally(filename)
    return public_url